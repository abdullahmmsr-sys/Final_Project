"""
Document Processor - Handles uploaded document parsing and chunking
"""
import os
import tempfile
from pathlib import Path
from typing import List, Dict, Any, BinaryIO
import aiofiles


from pypdf import PdfReader
from docx import Document as DocxDocument


class DocumentProcessor:
    
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    
    @staticmethod
    async def save_upload(file, filename: str) -> Path:
       
        try:
            from backend.config import UPLOAD_DIR
        except ImportError:
            from config import UPLOAD_DIR
        

        UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        
        suffix = Path(filename).suffix.lower()
        
        import uuid
        unique_name = f"{uuid.uuid4()}{suffix}"
        file_path = UPLOAD_DIR / unique_name
        
        
        content = await file.read()
        
       
        with open(file_path, 'wb') as f:
            f.write(content)
        
        return file_path
    
    @staticmethod
    def extract_text(file_path: Path) -> str:
       
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return DocumentProcessor._extract_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            return DocumentProcessor._extract_docx(file_path)
        elif suffix == '.txt':
            return DocumentProcessor._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    @staticmethod
    def _extract_pdf(file_path: Path) -> str:
        
        text_parts = []
        reader = PdfReader(str(file_path))
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def _extract_docx(file_path: Path) -> str:
        
        doc = DocxDocument(str(file_path))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def _extract_txt(file_path: Path) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def chunk_text(
        text: str, 
        chunk_size: int = 1000, 
        overlap: int = 100
    ) -> List[Dict[str, Any]]:

        chunks = []
        
        
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            
            if len(current_chunk) + len(para) > chunk_size and current_chunk:
                chunks.append({
                    "index": chunk_index,
                    "text": current_chunk.strip(),
                    "char_start": sum(len(c["text"]) for c in chunks),
                    "char_end": sum(len(c["text"]) for c in chunks) + len(current_chunk)
                })
                chunk_index += 1
                
                
                if overlap > 0 and len(current_chunk) > overlap:
                    current_chunk = current_chunk[-overlap:] + "\n\n" + para
                else:
                    current_chunk = para
            else:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
        
        
        if current_chunk.strip():
            chunks.append({
                "index": chunk_index,
                "text": current_chunk.strip(),
                "char_start": sum(len(c["text"]) for c in chunks),
                "char_end": sum(len(c["text"]) for c in chunks) + len(current_chunk)
            })
        
        return chunks
    
    @staticmethod
    def identify_sections(text: str) -> List[Dict[str, Any]]:
        """
        Identify policy/compliance sections in the document
        
        This helps map document content to specific control areas
        """
        sections = []
        lines = text.split('\n')
        
        current_section = None
        current_content = []
        
        
        section_indicators = [
            'policy', 'procedure', 'control', 'requirement',
            'standard', 'guideline', 'section', 'chapter',
            'سياسة', 'إجراء', 'ضابط', 'متطلب'  
        ]
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            
            is_header = (
                len(line) < 200 and
                len(line) > 3 and
                (
                    any(ind in line_lower for ind in section_indicators) or
                    line.isupper() or
                    (line.endswith(':') and len(line) < 100)
                )
            )
            
            if is_header:
               
                if current_section and current_content:
                    sections.append({
                        "title": current_section,
                        "content": "\n".join(current_content),
                        "line_start": sections[-1]["line_end"] + 1 if sections else 0
                    })
                
                current_section = line.strip()
                current_content = []
            else:
                if line.strip():
                    current_content.append(line)
        
        
        if current_section and current_content:
            sections.append({
                "title": current_section,
                "content": "\n".join(current_content),
                "line_start": sections[-1]["line_end"] + 1 if sections else 0,
                "line_end": len(lines)
            })
        
        return sections
    
    @staticmethod
    def cleanup(file_path: Path):
        """Remove temporary file"""
        try:
            os.unlink(file_path)
        except Exception:
            pass



document_processor = DocumentProcessor()
